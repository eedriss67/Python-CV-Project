# How to use python code to build a CV Resume

from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt
import pyttsx3


def intro():
    talk = pyttsx3.speak("This is my CV project. Thank you for checking it out.")
    return talk


intro()


def speak(message):
    speech = pyttsx3.speak(message)
    return speech


cv = Document()

# Adding style
font = cv.styles["Header"].font
font.size = Pt(18)
font.color.rgb = RGBColor(0, 0, 255)

style = cv.styles["Normal"]
font = style.font
font.name = "Ar Cena"
font.size = Pt(15)

cv.add_paragraph()

# Adding a profile picture
cv.add_picture("logo.png", width=Inches(1.5))

# Adding name, phone number and email
speak("Enter your full name")
name = input("Enter your full name: ")

speak("Enter your phone number")
phone_number = input("Enter your phone number: ")

speak("Enter your email address")
email = input("Enter your email address: ")

cv.add_paragraph(name + " | " + phone_number + " | " + email)

# Adding a heading
cv.add_heading("About me")

speak("Tell us about yourself")
about_me = input("Tell us about yourself: ")
cv.add_paragraph(about_me)

# Adding work experience heading
cv.add_heading("Work Experience")
info = cv.add_paragraph()

speak("Name of company you work at")
company = input("Name of company you work at: ")

speak("Date you started")
start_date = input("Date you started: ")

speak("Date you left")
end_date = input("Date you left: ")

info.add_run(company + " ").bold = True
info.add_run(start_date + " - " + end_date + "\n")

speak("Your position in the company")
position = input("Your position in the company: ")
info.add_run(position)

while True:
    speak("Do you have more work experience? Yes or No")
    more_experience = input("Do you have more work experience? Yes or No ")
    if more_experience.lower() == "yes":
        info = cv.add_paragraph()

        speak("Name of company you work at")
        company = input("Name of company you work at: ")

        speak("Date you started")
        start_date = input("Date you started: ")

        speak("Date you left")
        end_date = input("Date you left: ")

        info.add_run(company + " ").bold = True
        info.add_run(start_date + " - " + end_date + "\n")

        speak("Your position in the company")
        position = input("Your position in the company: ")
        info.add_run(position)
    else:
        break

# Adding Skill-set
cv.add_heading("Skills")
info = cv.add_paragraph()
info.style = "List Bullet"

speak("Enter your skill")
skill = input("Enter your skill ")
info.add_run(skill + "\n")

while True:
    speak("Do you have more skills? Yes or No")
    skill_set = input("Do you have more skills? Yes or No ")
    if skill_set.lower() == "yes":
        info = cv.add_paragraph()
        info.style = "List Bullet"

        speak("Enter other skill")
        skill = input("Enter other skill ")
        info.add_run(skill + "\n")
    else:
        break

# Adding a header
section = cv.sections[0]
header = section.header
info = header.paragraphs[0]
info.text = "MY CV RESUME"
info.alignment = 1

# Adding a footer
section = cv.sections[0]
footer = section.footer
info = footer.paragraphs[0]
info.text = "This CV was generated courtesy of Amigoscode YouTube Python Tutorial."

cv.save("cv.docx")
